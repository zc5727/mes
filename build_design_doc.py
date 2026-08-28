from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

OUT = Path('/Users/a1/Documents/ChatGPT/mes/MES智能制造运营平台设计与实施方案.docx')
IMG = Path('/Users/a1/Documents/ChatGPT/mes/MES智能制造平台架构预览.png')
BLUE = '17324D'; ACCENT='2E74B5'; LIGHT='EAF3FF'; GREEN='EEF9F3'; GOLD='FFF5E6'; PURPLE='F4EEFF'; MUTED='5B6B7A'; BORDER='CBD5E1'
FONT='HYZhongHeiKW'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, color=BORDER, sz='6'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:'+edge; el=borders.find(qn(tag))
        if el is None: el=OxmlElement(tag); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz); el.set(qn('w:color'),color)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run, size=10.5, color='1F2937', bold=False, italic=False):
    run.font.name=FONT; run._element.rPr.rFonts.set(qn('w:ascii'),FONT); run._element.rPr.rFonts.set(qn('w:hAnsi'),FONT); run._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def shade_para(p, fill, border=ACCENT):
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pPr.append(shd)
    pb=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8'); bot.set(qn('w:color'),border); pb.append(bot); pPr.append(pb)

def add_para(doc, text='', style=None, size=10.5, color='1F2937', bold=False, align=None, before=0, after=6, italic=False):
    p=doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    if align is not None: p.alignment=align
    r=p.add_run(text); set_run(r,size,color,bold,italic); return p

def add_bullets(doc, items, level=0):
    for item in items:
        p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
        p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1
        set_run(p.add_run(item),10.5)

def add_numbered(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1; set_run(p.add_run(item),10.5)

def add_table(doc, headers, rows, widths=None, header_fill='E8EEF5'):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; set_cell_shading(c,header_fill); set_cell_border(c); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); set_run(r,9.5,BLUE,True)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; set_cell_border(c); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(val)); set_run(r,9.2)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return table

def add_page_field(p):
    r=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.addnext(fld)

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); r=p.add_run(text); set_run(r,16 if level==1 else 13 if level==2 else 11.5, ACCENT if level<3 else BLUE, True); return p

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85); sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)
# styles
styles=doc.styles
for name,size,color,bold,space_before,space_after in [('Normal',10.5,'1F2937',False,0,6),('Heading 1',16,ACCENT,True,16,8),('Heading 2',13,ACCENT,True,12,6),('Heading 3',11.5,BLUE,True,8,4)]:
    st=styles[name]; st.font.name=FONT; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=bold; st.paragraph_format.space_before=Pt(space_before); st.paragraph_format.space_after=Pt(space_after); st.paragraph_format.line_spacing=1.15
# headers
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_run(header.add_run('MES智能制造运营平台  |  设计与实施方案'),8.5,MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run(footer.add_run('赵丞单人开发版  ·  内部设计文档  ·  '),8.5,MUTED); add_page_field(footer)
# cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run('制造数字化方案'); set_run(r,12,ACCENT,True)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(10); r=p.add_run('MES智能制造运营平台'); set_run(r,30,BLUE,True)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(26); r=p.add_run('需求分析、系统架构与实施方案'); set_run(r,18,MUTED,False)
# cover rule
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(20); shade_para(p,LIGHT,ACCENT); p.add_run(' ')
add_para(doc,'以设备互联、视觉理解、主动控制和对话决策为核心，构建面向生产预防与过程闭环的智能制造平台。',size=13,color=BLUE,bold=True,after=18)
meta=[('项目负责人','赵丞'),('开发模式','单人全栈开发'),('文档版本','V1.0'),('编制日期','2026年8月28日'),('文档状态','方案设计版')]
add_table(doc,['项目属性','内容'],meta,[1.55,4.9],header_fill='EAF3FF')
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(80); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; set_run(p.add_run('CONFIDENTIAL · 项目设计资料'),9,MUTED,True)
doc.add_page_break()
# toc/exec summary
heading(doc,'文档说明',1)
add_para(doc,'本文档用于指导MES智能制造运营平台的需求确认、架构设计、研发排期、人员规划、测试验证和试点实施。方案以“单人赵丞开发、先做单车间MVP、逐步扩展”为基本约束。',after=10)
heading(doc,'目录',2)
for x in ['1. 项目概述','2. 需求分析','3. 总体架构设计','4. 核心业务与主动控制策略','5. 数据与接口设计','6. 安全、权限与运维','7. 测试环境与质量保障','8. 人员配比与职责','9. 开发周期与实施计划','10. 验收标准与风险控制','附录A：一期范围清单','附录B：开源模块建议']: add_para(doc,x,size=10.5,after=2)
heading(doc,'执行摘要',2)
add_para(doc,'本项目不是单纯的生产报表系统，而是面向“生产中尽量不出问题”的主动式制造运营平台。系统通过设备、订单、物料、质量、图纸和现场视觉数据的融合，识别风险、提出方案、经过授权后执行，并对结果进行追踪和审计。',after=8)
add_para(doc,'首期不追求全功能，而是打通一条可验证闭环：模拟设备 → 生产工单 → 过程报工 → 异常预警 → 策略建议 → 厂长对话确认 → 派工/锁批次/通知 → 结果审计。',size=11,color=BLUE,bold=True,after=8)
doc.add_page_break()
# 1
heading(doc,'1. 项目概述',1)
heading(doc,'1.1 建设目标',2)
add_bullets(doc,['建立统一的生产、设备、质量、物料和图纸数据底座。','实现设备状态实时感知、工单过程可追踪、异常问题可闭环。','通过图纸和视觉分析，减少错图、错料、错工艺和漏检。','通过主动控制策略，提前识别瓶颈、故障、质量和交付风险。','通过厂长对话智能体，让管理者直接获得结论和决策建议。'])
heading(doc,'1.2 产品定位',2)
add_para(doc,'产品定位为“主动式智能制造运营平台”，而不是传统的菜单型MES。查询尽量对话化，决策结构化呈现，执行必须经过权限控制和审计。',after=8)
add_table(doc,['目标层','核心问题','平台能力'],[
('现场层','设备是否正常？是否会停？','设备接入、状态监控、报警、预测维护'),
('执行层','订单能否按期完成？','工单、排产、派工、报工、进度'),
('质量层','问题能否在扩散前发现？','首件、巡检、视觉、批次锁定、追溯'),
('管理层','厂长需要决策什么？','对话智能体、风险分析、方案模拟、审批执行')],[1.0,2.15,3.3])
heading(doc,'1.3 目标用户',2)
add_table(doc,['角色','主要任务','交互方式'],[
('厂长','掌握全局、处理重大风险、审批关键决策','对话、摘要卡片、方案确认'),
('生产主管','排产、派工、进度、异常协调','看板、工作台、智能助手'),
('设备主管','报警、维修、点检、保养','设备台账、工单、移动端'),
('质量主管','检验、放行、追溯、整改','质量工作台、图像复核'),
('操作员','接收任务、报工、检验、上报异常','平板/移动端、扫码、拍照、语音')],[1.0,2.9,2.55])
# 2
heading(doc,'2. 需求分析',1)
heading(doc,'2.1 业务主流程',2)
add_para(doc,'生产主流程：销售/生产订单 → 计划排产 → 工单拆解 → 物料准备 → 设备与人员确认 → 首件确认 → 批量生产 → 报工/检验 → 入库/交付 → 追溯与分析。',after=8)
add_para(doc,'异常主流程：设备报警或质量异常 → 风险分级 → 影响范围分析 → 生成策略 → 人工确认/自动执行 → 处理复核 → 解除锁定 → 效果评估 → 规则优化。',after=8)
heading(doc,'2.2 功能需求矩阵',2)
add_table(doc,['模块','一期需求','二期增强','优先级'],[
('基础数据','组织、用户、角色、设备、产品、BOM、工艺','多工厂、多租户、主数据同步','P0'),
('生产执行','订单、工单、排产、派工、报工、进度','智能排产、产能仿真','P0'),
('设备管理','MQTT/OPC UA、状态、报警、点检、维修','预测性维护、备件优化','P0'),
('质量管理','首件、巡检、不良、批次锁定、追溯','视觉缺陷模型、SPC、8D知识库','P0'),
('图纸视觉','上传、版本、OCR、字段提取、人工确认','CAD深度解析、版本差异比对','P1'),
('智能表单','报工、点检、检验、异常、拍照、语音','低代码表单设计器','P1'),
('对话智能体','生产/设备/质量查询、日报、待办','跨域方案模拟、语音交互','P1'),
('数字孪生','二维车间、设备状态、订单下钻','三维工厂、物流仿真','P2')],[1.0,2.3,2.5,0.65])
heading(doc,'2.3 非功能需求',2)
add_table(doc,['类别','要求'],[
('可靠性','关键业务可恢复；设备数据支持断点续传和重复消息去重。'),
('实时性','设备状态通常5秒内刷新；报警消息优先传递；看板支持WebSocket。'),
('可审计','图纸、工艺、批次、停线、放行、派工和智能体动作全部留痕。'),
('安全性','统一身份认证、角色权限、数据权限、最小权限和敏感文件隔离。'),
('可扩展','设备协议、表单、规则、看板和智能体工具采用插件或配置化方式。'),
('可维护','模块边界清晰；核心制造规则不写死在页面和大模型提示词中。')],[1.2,5.25])
doc.add_page_break()
# arch
heading(doc,'3. 总体架构设计',1)
add_para(doc,'采用“边缘接入 + 数据平台 + MES业务 + 主动控制 + 对话体验”的分层架构。单人开发阶段优先采用模块化单体，设备接入和视觉服务独立部署，避免过早引入复杂微服务。',after=8)
if IMG.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(IMG),width=Inches(6.55))
    add_para(doc,'图1  MES智能制造运营平台总体架构',size=9,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
heading(doc,'3.1 分层说明',2)
add_table(doc,['层级','主要组件','职责'],[
('体验与决策层','厂长智能体、看板、表单、移动端、数字孪生','按角色提供结论、待办和操作入口。'),
('业务应用层','生产、设备、质量、物料、图纸视觉、审批通知','承载MES业务对象和标准业务流程。'),
('主动控制层','状态识别、规则、预测、模拟、优化、执行、审计','判断风险并生成可解释、可确认的控制策略。'),
('数据与集成层','API、权限、消息总线、业务库、时序库、文件库','统一数据访问、事件通信、存储和审计。'),
('边缘现场层','网关、PLC、CNC、机器人、传感器、相机、虚拟产线','连接现场并提供实时采集、控制和模拟。')],[1.25,2.7,2.5])
heading(doc,'3.2 推荐技术栈',2)
add_table(doc,['领域','技术选型','说明'],[
('前端','Vue 3 + TypeScript + ECharts + Three.js','MES工作台、看板、孪生和对话界面。'),
('业务后端','NestJS + TypeScript + Prisma','生产、质量、设备、权限和审计。'),
('设备/视觉','Python + FastAPI + OpenCV','协议接入、模拟器、OCR和视觉处理。'),
('数据','PostgreSQL + TimescaleDB + Redis','业务数据、时序数据和缓存。'),
('消息','MQTT + Kafka/NATS + WebSocket','设备上行、业务事件和实时推送。'),
('部署','Docker Compose + MinIO + Prometheus/Loki','本地开发、试点部署和基础运维。')],[1.25,2.6,2.6])
# 4
heading(doc,'4. 核心业务与主动控制策略',1)
heading(doc,'4.1 主动控制策略框架',2)
add_para(doc,'主动控制策略必须由规则、约束、预测和方案模拟共同组成。大模型只负责理解自然语言、解释结果和调用工具，不直接决定高风险生产动作。',after=8)
add_numbered(doc,['实时感知：采集设备、订单、物料、质量、人员和现场视觉数据。','状态识别：判断正常、预警、异常、瓶颈、离线等状态。','风险预测：计算故障、质量、拥堵、物料和延期风险。','方案生成：生成调度、备用设备、降速、停线、锁批次等候选方案。','约束校验：校验设备能力、工艺、物料、人员资质和安全边界。','影响评估：比较交期、产能、成本、质量和风险影响。','确认执行：低风险自动执行，中高风险对话审批，高风险禁止自动执行。','效果评估：对比策略执行前后结果，沉淀规则和经验。'])
heading(doc,'4.2 首批策略清单',2)
add_table(doc,['策略编号','触发条件','系统动作','权限级别'],[
('STR-001','设备温度/振动持续进入预警区','通知设备主管，创建点检任务','自动'),
('STR-002','连续3件不良或关键尺寸超限','暂停工序，锁定当前批次','确认后执行'),
('STR-003','前工序在制品超过缓冲上限','降低前工序节拍，提示生产主管','确认后执行'),
('STR-004','关键设备故障影响订单','推荐备用设备并模拟交期影响','确认后执行'),
('STR-005','物料库存低于安全时长','预警仓库和计划人员','自动'),
('STR-006','工艺/图纸/物料不匹配','禁止开工并生成核验任务','硬约束'),
('STR-007','订单延期风险超过阈值','升级通知并生成调整方案','确认后执行')],[0.8,2.0,2.65,1.0])
heading(doc,'4.3 厂长智能体交互规范',2)
add_para(doc,'厂长端只提供“问工厂、看重点、处理待办”三个主要入口。系统回答必须包含结论、依据、影响、建议和可执行动作。',after=6)
add_para(doc,'示例：二号产线预计45分钟后形成瓶颈。原因是CNC-03节拍下降12%，前工序来料持续增加。建议将订单A的30%任务切换到CNC-05，预计避免延期2小时。是否确认执行？',size=10.5,color=BLUE,bold=True,after=8)
# 5
heading(doc,'5. 数据与接口设计',1)
heading(doc,'5.1 核心数据对象',2)
add_table(doc,['对象','关键字段','关联对象'],[
('设备','设备编码、类型、状态、能力、参数、协议','产线、工艺、报警、维修、工单'),
('产品/图纸','产品编码、图号、版本、材料、尺寸、公差','BOM、工艺、检验、订单'),
('生产订单','订单号、产品、数量、交期、优先级','工单、物料、批次、报工'),
('工单/工序','工单号、设备、人员、状态、计划时间','订单、工艺、设备、质量'),
('质量记录','检验项目、标准、实测值、结果、图片','工单、批次、图纸、人员'),
('异常事件','类型、等级、来源、影响、处理、状态','设备、订单、批次、策略'),
('策略决策','触发条件、候选方案、依据、审批、执行结果','风险、工单、设备、审计')],[1.15,2.35,2.95])
heading(doc,'5.2 设备消息规范',2)
add_para(doc,'统一设备事件格式，建议至少包含设备标识、点位、数值、状态、采集时间、网关时间、质量码和关联批次。消息需要支持幂等键，避免断线重传造成重复数据。',after=6)
add_para(doc,'示例：device_id、metric、value、unit、status、event_time、gateway_id、quality、trace_id。',size=10.5,color=BLUE,after=8)
heading(doc,'5.3 文件和图纸处理流程',2)
add_numbered(doc,['上传原始图纸并生成文件指纹。','识别图号、版本、标题栏、表格、尺寸、公差和技术要求。','将识别结果存为待确认草稿，显示置信度和原图定位。','由工程或质量人员人工确认。','确认后生成BOM、工艺建议、检验项目或质量控制计划。','将正式版本绑定到生产订单，历史版本只读归档。'])
# 6
heading(doc,'6. 安全、权限与运维',1)
heading(doc,'6.1 权限体系',2)
add_bullets(doc,['统一账号体系：用户、组织、岗位、角色和权限集中管理。','数据权限：按工厂、车间、产线、设备、产品和订单范围控制。','操作权限：查询、编辑、审批、停线、放行、派工分别授权。','文件权限：图纸、工艺、质量附件和现场图片按部门和角色隔离。','智能体权限：每次工具调用都必须携带用户身份、角色和业务范围。'])
heading(doc,'6.2 审计要求',2)
add_bullets(doc,['图纸上传、解析、确认、发布和版本变更。','工艺参数、检验结果、批次锁定和解除。','停线、放行、派工、排产和策略执行。','厂长智能体的查询、建议、审批和实际执行动作。','设备数据异常、断线、补传和人工修正。'])
heading(doc,'6.3 运维要求',2)
add_table(doc,['项目','要求'],[
('备份','业务库每日备份；文件库按版本和对象存储策略备份。'),
('监控','服务健康、消息堆积、设备在线率、接口错误、策略失败率。'),
('恢复','关键服务支持重启恢复；设备断线后支持补传和幂等处理。'),
('部署','开发、测试、试点环境使用独立配置和独立数据。'),
('升级','数据库迁移、策略版本、模型版本和前端版本可追踪。')],[1.1,5.35])
# 7
heading(doc,'7. 测试环境与质量保障',1)
heading(doc,'7.1 无真实工厂环境下的模拟方案',2)
add_para(doc,'采用虚拟产线替代真实PLC和设备，先验证数据链路、业务流程和主动控制策略。',after=6)
add_table(doc,['模拟对象','示例'],[
('产线','原料仓 → CNC-01 → CNC-02 → 检测工位 → 成品仓'),
('设备数据','温度、振动、电流、转速、节拍、产量、良率、报警、离线'),
('生产数据','订单、工单、计划数量、实际数量、在制品、交期'),
('故障注入','过热、振动升高、设备离线、连续不良、物料不足、网络中断'),
('视觉素材','正常图纸、低清图纸、多版本图纸、纸质表单、缺陷图片')],[1.5,4.95])
heading(doc,'7.2 测试分层',2)
add_bullets(doc,['单元测试：规则、权限、状态机、数据校验、策略计算。','接口测试：设备消息、业务API、文件解析、智能体工具调用。','集成测试：设备模拟器 → MQTT → MES → 看板 → 策略引擎。','场景测试：设备故障、质量连续不良、产线拥堵、订单延期。','恢复测试：断网、重复消息、服务重启、数据补传。','用户验收：厂长、生产、设备、质量和操作员分别验证。'])
heading(doc,'7.3 关键质量门禁',2)
add_bullets(doc,['AI识别结果必须人工确认后才能成为正式业务数据。','高风险动作不能由大模型直接执行。','异常批次未解除锁定前不能入库或继续流转。','图纸和工艺版本必须与生产订单固定绑定。','所有关键动作必须能够追溯到人、时间、依据和结果。'])
# 8
heading(doc,'8. 人员配比与职责',1)
heading(doc,'8.1 赵丞单人开发模式',2)
add_para(doc,'赵丞负责产品、架构、前后端、数据库、设备模拟、AI集成和部署。为了控制范围，建议采用模块化单体，不在第一阶段自行训练视觉模型，也不开发复杂三维引擎。',after=8)
add_table(doc,['职责领域','单人承担内容','建议外部支持'],[
('产品与业务','需求、流程、原型、策略规则、验收','现场工艺/质量专家评审'),
('前端','MES工作台、看板、对话、表单、孪生','UI视觉评审'),
('后端','订单、工单、质量、设备、权限、审计','安全评审'),
('设备接入','MQTT、OPC UA、Modbus、模拟器','真实设备厂商配合'),
('AI视觉','OCR、图纸字段、图片分析、Agent工具','视觉算法或模型顾问'),
('测试运维','自动化测试、部署、备份、监控','试点现场人员参与')],[1.25,3.0,2.2])
heading(doc,'8.2 标准团队配置',2)
add_table(doc,['角色','人数','投入重点'],[
('产品/制造顾问','1','流程、工艺、质量规则和验收'),('技术负责人','1','架构、安全、数据和关键代码'),('前端','1','业务、看板、智能体和移动端'),('后端','2','MES、权限、事件和策略服务'),('设备/边缘','1','协议、网关、现场数据'),('AI/视觉','1','图纸、OCR、缺陷和模型服务'),('测试/实施','1','测试、培训、试点和数据治理')],[1.55,0.7,4.2])
# 9
heading(doc,'9. 开发周期与实施计划',1)
heading(doc,'9.1 单人开发计划',2)
add_table(doc,['阶段','周期','主要交付物'],[
('需求与原型','3-4周','需求说明、流程、原型、数据字典、策略清单'),('基础平台','6-8周','权限、组织、设备台账、文件中心、审计'),('MES核心','8-12周','订单、工单、排产、报工、质量、追溯'),('设备与模拟器','6-10周并行','MQTT、OPC UA、设备状态、故障注入'),('视觉与表单','6-10周','图纸解析、人工确认、拍照/语音填报'),('智能体与策略','8-12周','对话查询、预警、方案模拟、审批执行'),('试点上线','6-8周','单车间试点、培训、修复、验收')],[1.35,1.25,3.85])
add_para(doc,'预计单人完成稳定的单车间试点版本需要12-18个月；如果只做可演示MVP，可压缩到4-6个月，但不能等同于生产级系统。',size=11,color=BLUE,bold=True,after=8)
heading(doc,'9.2 推荐里程碑',2)
add_table(doc,['里程碑','完成标准'],[
('M1 设备沙盒','虚拟设备可持续发送数据，支持故障注入和恢复。'),
('M2 MES闭环','订单、工单、报工、质量和批次追溯可完整运行。'),
('M3 主动预警','至少7个策略可触发、解释、审批和审计。'),
('M4 视觉闭环','图纸上传可提取字段，确认后可生成检验项目。'),
('M5 厂长体验','厂长通过对话获得生产结论并处理待办。'),
('M6 单车间试点','真实业务连续运行，关键数据和异常闭环可验证。')],[1.6,4.85])
# 10
heading(doc,'10. 验收标准与风险控制',1)
heading(doc,'10.1 一期验收标准',2)
add_bullets(doc,['支持一个工厂、一个车间、一条产线、三类设备和两类产品。','设备数据稳定接入，状态和报警能够在看板中展示。','生产订单能够拆解为工单并完成派工和报工。','图纸能够上传、解析、人工确认并绑定产品/工单。','首件和巡检结果能够关联图纸、批次、设备和人员。','异常能够自动分级，支持通知、锁批次、停工和处理复核。','厂长可以通过对话查询生产、设备和质量情况。','系统可以生成至少两个候选策略并展示影响评估。','关键操作、AI调用和审批均有审计记录。'])
heading(doc,'10.2 主要风险',2)
add_table(doc,['风险','影响','应对措施'],[
('范围失控','周期和维护成本快速上升','严格采用P0/P1/P2，首期只做单车间'),('真实数据缺失','预测和策略无法验证','先建设虚拟产线和故障注入'),('设备协议复杂','接入周期不可控','先统一MQTT，OPC/Modbus选1-2种'),('视觉识别误差','产生错检或错放行','AI草稿+人工确认，不直接放行'),('策略误动作','影响生产和安全','硬约束、审批、审计、回滚'),('单人瓶颈','开发和实施互相挤压','优先复用开源模块，减少自研外围能力')],[1.25,2.2,3.0])
# appendix
heading(doc,'附录A：一期范围清单',1)
add_table(doc,['范围内','暂不纳入'],[
('单车间设备接入','全厂多工厂协同'),('订单/工单/报工','复杂高级APS排产'),('设备状态/报警/点检','全自动预测性维护'),('首件/巡检/批次锁定','高精度全品类视觉检测'),('图纸OCR与人工确认','完整CAD几何语义解析'),('厂长对话查询与审批','大规模多智能体自治'),('二维/轻量孪生','复杂三维仿真和数字工厂')],[3.2,3.25])
heading(doc,'附录B：开源模块建议',1)
add_table(doc,['模块','建议用途','集成方式'],[
('Node-RED','虚拟设备、流程编排、故障注入','MQTT/HTTP接入MES'),('DGIOT','工业协议、设备模型、告警、基础孪生','API/MQTT事件集成'),('MinerU','PDF、扫描文档、表格和图纸文本解析','Python服务封装'),('OpenCV','图像预处理和基础视觉','Python视觉服务'),('OpenMES','参考MES业务对象和流程','参考或局部二开，不建议直接承载全部策略'),('n8n','外围通知、审批和智能体工作流','Webhook/API调用')],[1.25,3.0,2.2])
add_para(doc,'注：开源模块实际用于商业部署前，应逐项核对许可证、版本、依赖和安全状况。',size=9,color=MUTED,italic=True,after=6)
# final signoff
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(25); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_run(p.add_run('方案负责人：赵丞\n开发模式：单人全栈开发\n'),10.5,BLUE,True)
# core properties
props=doc.core_properties; props.title='MES智能制造运营平台设计与实施方案'; props.subject='需求分析、系统架构、人员配比与开发周期'; props.author='赵丞'; props.keywords='MES, 智能制造, 主动控制, 数字孪生, 视觉分析, 智能体'
doc.save(OUT)
print(OUT)
