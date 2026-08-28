from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root=Path('/Users/a1/Documents/ChatGPT/mes')
font='/Users/a1/Library/Fonts/HYZhongHeiKW.ttf'
def f(n): return ImageFont.truetype(font,n)
# SaaS architecture image
W,H=1800,1100
im=Image.new('RGB',(W,H),'#f7f9fc'); d=ImageDraw.Draw(im)
def box(x,y,w,h,title,sub,color,fill='#ffffff'):
 d.rounded_rectangle((x,y,x+w,y+h),16,fill=fill,outline=color,width=4)
 d.text((x+18,y+15),title,font=f(26),fill='#17324d'); d.text((x+18,y+56),sub,font=f(17),fill='#526579')
def arr(a,b,color='#73849a'):
 d.line((*a,*b),fill=color,width=5); x2,y2=b; x1,y1=a
 import math
 ang=math.atan2(y2-y1,x2-x1); L=16
 p1=(x2-L*math.cos(ang-.45),y2-L*math.sin(ang-.45)); p2=(x2-L*math.cos(ang+.45),y2-L*math.sin(ang+.45))
 d.polygon([(x2,y2),p1,p2],fill=color)
d.text((45,25),'SaaS云边协同架构',font=f(36),fill='#17324d')
box(80,120,500,130,'SaaS控制平面','租户 / 订阅 / 配额 / 版本 / 运维','#4a90e2','#eaf3ff')
box(650,120,500,130,'统一身份与权限中心','SSO / RBAC / 数据权限 / 审计','#4a90e2','#eaf3ff')
box(1220,120,500,130,'厂长对话与运营门户','看板 / 智能体 / 待办 / 方案审批','#4a90e2','#eaf3ff')
for x,t,s in [(80,'生产与设备服务','订单 / 工单 / 设备 / 质量'),(420,'图纸视觉服务','文件 / OCR / 图片 / 表单'),(760,'主动策略服务','规则 / 预测 / 模拟 / 执行'),(1100,'数据平台','业务库 / 时序库 / 文件库'),(1440,'通知与集成','Webhook / 企业微信 / API')]: box(x,370,280,110,t,s,'#37a169','#eef9f3')
box(80,650,480,150,'租户A边缘网关','现场缓存 / 协议接入 / 本地规则 / 断网续传','#d64545','#fff0f0')
box(660,650,480,150,'租户B边缘网关','现场缓存 / 协议接入 / 本地规则 / 断网续传','#d64545','#fff0f0')
box(1240,650,480,150,'租户N边缘网关','现场缓存 / 协议接入 / 本地规则 / 断网续传','#d64545','#fff0f0')
for x in (330,900,1480): arr((x,250),(x,370),'#4a90e2')
for x in (320,900,1480): arr((x,480),(x,650),'#37a169')
for x in (210,790,1370):
 box(x,900,260,80,'PLC / CNC / 机器人','传感器 / 相机 / AGV','#d64545','#ffffff'); arr((x+130,800),(x+130,900),'#d64545')
img=root/'SaaS云边协同架构预览.png'; im.save(img)

path=root/'MES智能制造运营平台设计与实施方案.docx'; doc=Document(path)
# Add cover metadata row to first table
for tbl in doc.tables[:1]:
 row=tbl.add_row().cells
 row[0].text='交付形态'; row[1].text='多租户SaaS（云边协同）'
 break
# append SaaS section
p=doc.add_paragraph(); p.add_run().add_break()
doc.add_page_break()
def h(text,level=1):
 p=doc.add_paragraph(style=f'Heading {level}'); p.add_run(text); return p
def para(text,bold=False,color=None):
 p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
 r=p.add_run(text); r.bold=bold; r.font.name='HYZhongHeiKW'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'HYZhongHeiKW'); r.font.size=Pt(10.5)
 if color: r.font.color.rgb=__import__('docx').shared.RGBColor.from_string(color)
 return p
def table(headers,rows):
 t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
 for i,v in enumerate(headers):
  c=t.rows[0].cells[i]; c.text=v
  for r in c.paragraphs[0].runs: r.bold=True; r.font.name='HYZhongHeiKW'; r.font.size=Pt(9.5)
  shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'E8EEF5'); c._tc.get_or_add_tcPr().append(shd)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row):
   cells[i].text=str(v)
   for r in cells[i].paragraphs[0].runs: r.font.name='HYZhongHeiKW'; r.font.size=Pt(9.2)
 return t
h('11. SaaS交付架构补充',1)
para('本项目最终以多租户SaaS形式交付，采用“云端控制平面 + 企业租户空间 + 工厂边缘网关”的云边协同模式。SaaS平台负责业务、分析、智能体和策略编排，边缘网关负责设备接入、本地缓存和必要的现场规则。')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(img),width=Inches(6.5))
para('图2  SaaS云边协同架构',color='5B6B7A')
h('11.1 SaaS核心能力',2)
table(['能力域','需求内容'],[
('多租户','租户注册、企业管理员、工厂/车间/产线、租户配置、租户级看板。'),
('订阅与配额','套餐、设备数量、用户数量、数据点、AI调用量、文件存储量和到期控制。'),
('统一身份','SSO、账号、角色、岗位、组织和数据权限。'),
('租户隔离','业务库、时序数据、文件、图纸、图片、知识库和Agent会话隔离。'),
('边缘接入','按工厂部署网关，支持OPC UA、MQTT、Modbus、断网续传和本地缓存。'),
('运营运维','租户开通、版本发布、服务监控、日志、备份、恢复和用量统计。')])
h('11.2 云边协同边界',2)
table(['云端负责','边缘负责'],[
('租户、订阅、权限、业务数据','现场协议和设备连接'),('MES订单、工单、质量和追溯','本地采集、缓存和断点续传'),('策略生成、方案模拟和审批','低延迟报警和现场必要规则'),('厂长智能体和经营分析','断网时的基本生产可用性'),('模型、规则和配置版本管理','设备侧执行结果回传')])
h('11.3 SaaS安全与隔离要求',2)
for txt in ['每条业务数据必须带tenant_id，并同时记录工厂、车间、产线和用户范围。','不同租户禁止互相查询、检索、读取附件或调用对方Agent上下文。','图纸、表单图片、视觉结果和知识库按租户分区存储。','智能体工具调用必须携带用户身份、租户、角色、session_id和trace_id。','高风险生产动作必须由授权用户确认，云端不能绕过边缘安全边界直接控制设备。','租户级备份、恢复、导出和删除操作必须可审计。']:
 p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(txt)
h('11.4 SaaS套餐建议',2)
table(['版本','主要能力','建议计费维度'],[
('基础版','设备管理、工单、报工、基础看板','工厂/设备/用户'),('专业版','质量、追溯、图纸解析、智能表单','工厂/设备/文件/用户'),('智能版','厂长智能体、主动控制、风险预测','AI调用/策略/设备'),('企业版','私有化、专属部署、定制接口、专属服务','项目制/服务合同')])
h('11.5 SaaS实施周期调整',2)
para('由于增加多租户、订阅配额、云边协同、租户隔离、云端运维和商业化运营能力，单人开发周期应按12-18个月完成稳定单车间试点版本进行规划；SaaS商业MVP建议18-24个月，稳定商业平台建议24-36个月。')
# update last metadata
props=doc.core_properties; props.subject='SaaS多租户MES需求分析、系统架构、人员配比与实施方案'; props.author='赵丞'
doc.save(path)
print(path)
